#coding: utf-8

from PIL import Image
from collections import Counter
import docx
import os
import Spider_script
import plotWordcloud

def image_process(image):
    im = Image.open(image)
    pix = im.load()
    width = im.size[0]
    height = im.size[1]
    for x in range(width):
        for y in range(height):
            r,g,b = pix[x,y]
            if r <= 200 or g<= 200 or b<=200:
                pix[x,y] = (0,0,0)

    im.save('plotWordcloud/Images/bit_icon2.jpeg')

def write_file_with_file(r_file, w_file):
    file = docx.Document(r_file)
    shudong = file.open(w_file, 'w')
    for para in file.paragraphs:
        shudong.write(para)

    print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出每一段的内容
    for para in file.paragraphs:
        print(para.text)

    # 输出段落编号及段落内容
    for i in range(len(file.paragraphs)):
        print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)

def file_combine(files, lode_to):
        path = files
        files = os.listdir(path)
        with open(lode_to, 'w', encoding='utf-8') as fw:
            for filename in files:
                filedir = path + filename
                with open(filedir, encoding='utf-8') as fr:
                    for line in fr:
                        # print(line[0:1])
                        fw.write(line)

def write_file_with_list(lst, w_file):
    with open(w_file, 'w') as wf:
        for i in lst:
            wf.write(i)

if __name__ == '__main__':
    # 文件封装

    restrict_word2 = ['不','好','在','是','有','吗','就','想','个','的','我','一个','很','都','上']
    # wordcloud_generate('doc/十九大报告全文.txt', 'plotWordcloud/Images/alice_mask.png', 'userdict/userdict.txt', 'plotWordcloud/Images/十九大.png',restrict_word2)
    # wordcloud_generate('doc/树洞.txt', 'plotWordcloud/Images/alice_mask.png', 'userdict/userdict.txt', 'plotWordcloud/Images/树洞.png', restrict_word2)
    # plw.wordcloud_generate('/Users/mac/PycharmProjects/demoWorldcloud/doc/二手兼职.txt','Images/music.png','/Users/mac/PycharmProjects/demoWorldcloud/userdict/userdict.txt','Images/二手.png',restrict_word2)
    # file_combine('doc/二手兼职/','doc/二手兼职2.txt')
    write_file_with_list(Spider_script.Spider("", 10), 'doc/贴吧/data.txt')
    plotWordcloud.wordcloud_generate('/Users/mac/PycharmProjects/demoWorldcloud/doc/贴吧/data.txt','Images/music.png','/Users/mac/PycharmProjects/demoWorldcloud/userdict/userdict.txt','Images/贴吧.png',restrict_word2)


